#!/usr/bin/env python3
"""
Export CDA Dubai Pitch to Word Document with Arabic RTL formatting
"""

import json
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import os

# Colors
MAROON = RGBColor(0xC4, 0x1E, 0x3A)
GOLD = RGBColor(0xD4, 0xAF, 0x37)
DARK_NAVY = RGBColor(0x1A, 0x1A, 0x2E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY = RGBColor(0x66, 0x66, 0x66)


def set_rtl_paragraph(paragraph):
    """Set paragraph to RTL direction"""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_rtl_run(run):
    """Set run to RTL direction"""
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)


def add_heading_rtl(doc, text, level=1, color=MAROON):
    """Add RTL heading"""
    heading = doc.add_heading(level=level)
    set_rtl_paragraph(heading)
    run = heading.add_run(text)
    run.font.color.rgb = color
    run.font.name = 'Tajawal'
    run._element.rPr.rFonts.set(qn('w:cs'), 'Tajawal')
    set_rtl_run(run)
    return heading


def add_paragraph_rtl(doc, text, bold=False, color=None, size=12):
    """Add RTL paragraph"""
    para = doc.add_paragraph()
    set_rtl_paragraph(para)
    run = para.add_run(text)
    run.font.name = 'Tajawal'
    run._element.rPr.rFonts.set(qn('w:cs'), 'Tajawal')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    set_rtl_run(run)
    return para


def add_bullet_rtl(doc, text, level=0):
    """Add RTL bullet point"""
    para = doc.add_paragraph(style='List Bullet')
    set_rtl_paragraph(para)
    para.paragraph_format.left_indent = Cm(level * 0.5)
    run = para.add_run(text)
    run.font.name = 'Tajawal'
    run._element.rPr.rFonts.set(qn('w:cs'), 'Tajawal')
    run.font.size = Pt(11)
    set_rtl_run(run)
    return para


def set_table_rtl(table):
    """Set table to RTL"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    bidiVisual = OxmlElement('w:bidiVisual')
    tblPr.append(bidiVisual)


def create_table_rtl(doc, headers, rows, header_color=MAROON):
    """Create RTL table with headers"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    set_table_rtl(table)
    
    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        para = cell.paragraphs[0]
        set_rtl_paragraph(para)
        run = para.add_run(header)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.name = 'Tajawal'
        run._element.rPr.rFonts.set(qn('w:cs'), 'Tajawal')
        set_rtl_run(run)
        # Set background color
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'C41E3A')
        cell._tc.get_or_add_tcPr().append(shading)
    
    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            cell = row_cells[col_idx]
            para = cell.paragraphs[0]
            set_rtl_paragraph(para)
            run = para.add_run(str(cell_text))
            run.font.name = 'Tajawal'
            run._element.rPr.rFonts.set(qn('w:cs'), 'Tajawal')
            run.font.size = Pt(10)
            set_rtl_run(run)
    
    return table


def add_page_break(doc):
    """Add page break"""
    doc.add_page_break()


def render_cover(doc, slide):
    """Render cover slide"""
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Title
    title = add_heading_rtl(doc, slide['title'], level=0, color=MAROON)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    if 'subtitle' in slide:
        sub = add_paragraph_rtl(doc, slide['subtitle'], bold=True, size=16, color=GOLD)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Tagline
    if 'tagline' in slide:
        tag = add_paragraph_rtl(doc, slide['tagline'], size=14, color=GRAY)
        tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Year
    if 'year' in slide:
        year = add_paragraph_rtl(doc, slide['year'], size=14, color=MAROON)
        year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_page_break(doc)


def render_executive_summary(doc, slide):
    """Render executive summary"""
    add_heading_rtl(doc, slide['title'], level=1)
    
    if 'vision' in slide:
        add_paragraph_rtl(doc, "الرؤية:", bold=True, color=MAROON, size=13)
        add_paragraph_rtl(doc, slide['vision'], size=12)
    
    if 'key_points' in slide:
        doc.add_paragraph()
        add_paragraph_rtl(doc, "النقاط الرئيسية:", bold=True, color=MAROON, size=13)
        for point in slide['key_points']:
            add_bullet_rtl(doc, point)
    
    if 'expected_outcomes' in slide:
        doc.add_paragraph()
        add_paragraph_rtl(doc, "النتائج المتوقعة:", bold=True, color=MAROON, size=13)
        for outcome in slide['expected_outcomes']:
            add_bullet_rtl(doc, outcome)
    
    add_page_break(doc)


def render_framework_overview(doc, slide):
    """Render framework overview"""
    add_heading_rtl(doc, slide['title'], level=1)
    add_paragraph_rtl(doc, slide.get('description', ''), size=12)
    doc.add_paragraph()
    
    if 'tracks' in slide:
        for track in slide['tracks']:
            icon = track.get('icon', '')
            name = track.get('name', '')
            add_paragraph_rtl(doc, f"{icon} {name}", bold=True, color=MAROON, size=13)
            add_paragraph_rtl(doc, track.get('description', ''), size=11, color=GRAY)
            if 'components' in track:
                for comp in track['components']:
                    add_bullet_rtl(doc, comp)
            doc.add_paragraph()
    
    add_page_break(doc)


def render_section_divider(doc, slide):
    """Render section divider"""
    doc.add_paragraph()
    doc.add_paragraph()
    
    icon = slide.get('icon', '')
    title = slide.get('title', '')
    subtitle = slide.get('subtitle', '')
    
    t = add_heading_rtl(doc, f"{icon} {title}", level=1, color=MAROON)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if subtitle:
        s = add_paragraph_rtl(doc, subtitle, bold=True, size=18, color=GOLD)
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if 'description' in slide:
        d = add_paragraph_rtl(doc, slide['description'], size=12, color=GRAY)
        d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_page_break(doc)


def render_pillar_detail(doc, slide):
    """Render pillar detail - handles multiple content types"""
    add_heading_rtl(doc, slide['title'], level=1)
    
    if 'description' in slide:
        add_paragraph_rtl(doc, slide['description'], size=11, color=GRAY)
    doc.add_paragraph()
    
    content_type = slide.get('content_type', '')
    
    # Narrative Architecture
    if content_type == 'narrative_architecture' and 'narrative_architecture' in slide:
        na = slide['narrative_architecture']
        if 'core_narrative' in na:
            add_paragraph_rtl(doc, "السردية الأساسية:", bold=True, color=MAROON)
            add_paragraph_rtl(doc, na['core_narrative'], size=12)
        
        if 'narrative_pillars' in na:
            doc.add_paragraph()
            add_paragraph_rtl(doc, "ركائز السردية:", bold=True, color=MAROON)
            for pillar in na['narrative_pillars']:
                add_paragraph_rtl(doc, f"• {pillar.get('pillar', '')}", bold=True, size=11)
                add_paragraph_rtl(doc, pillar.get('message', ''), size=11)
                if 'proof_points' in pillar:
                    for pp in pillar['proof_points']:
                        add_bullet_rtl(doc, pp, level=1)
        
        if 'tone_of_voice' in na:
            doc.add_paragraph()
            add_paragraph_rtl(doc, "نبرة الصوت:", bold=True, color=MAROON)
            tov = na['tone_of_voice']
            add_paragraph_rtl(doc, f"الأساسية: {tov.get('primary', '')}", size=11)
            add_paragraph_rtl(doc, f"الثانوية: {tov.get('secondary', '')}", size=11)
    
    # Stakeholder Map
    elif content_type == 'stakeholder_map' and 'stakeholder_map' in slide:
        sm = slide['stakeholder_map']
        if 'circles' in sm:
            for circle in sm['circles']:
                add_paragraph_rtl(doc, circle.get('name', ''), bold=True, color=MAROON)
                if 'stakeholders' in circle:
                    headers = ['الفئة', 'العدد', 'الأولوية']
                    rows = [[s.get('name', ''), s.get('count', ''), s.get('priority', '')] 
                            for s in circle['stakeholders']]
                    create_table_rtl(doc, headers, rows)
                    doc.add_paragraph()
    
    # Segmentation Approach
    elif content_type == 'segmentation_approach' and 'segmentation_approach' in slide:
        sa = slide['segmentation_approach']
        if 'methodology' in sa:
            add_paragraph_rtl(doc, f"المنهجية: {sa['methodology']}", size=11)
        
        if 'segments' in sa:
            doc.add_paragraph()
            for seg in sa['segments']:
                add_paragraph_rtl(doc, f"• {seg.get('name', '')} ({seg.get('size', '')})", bold=True, size=11)
                add_paragraph_rtl(doc, seg.get('characteristics', ''), size=10, color=GRAY)
                if 'key_needs' in seg:
                    add_paragraph_rtl(doc, f"الاحتياجات: {' | '.join(seg['key_needs'])}", size=10)
                if 'channels' in seg:
                    add_paragraph_rtl(doc, f"القنوات: {' | '.join(seg['channels'])}", size=10)
                doc.add_paragraph()
    
    # Measurement Framework
    elif content_type == 'measurement_framework' and 'measurement_framework' in slide:
        mf = slide['measurement_framework']
        if 'model' in mf:
            add_paragraph_rtl(doc, f"النموذج: {mf['model']}", bold=True)
        
        if 'dimensions' in mf:
            doc.add_paragraph()
            headers = ['البُعد', 'الوزن', 'الحالي', 'السنة 1', 'السنة 3']
            rows = [[d.get('dimension', ''), d.get('weight', ''), d.get('current', ''),
                    d.get('target_y1', ''), d.get('target_y3', '')] for d in mf['dimensions']]
            create_table_rtl(doc, headers, rows)
        
        if 'overall_score' in mf:
            doc.add_paragraph()
            os = mf['overall_score']
            add_paragraph_rtl(doc, f"المؤشر الإجمالي: {os.get('current', '')} → {os.get('year1', '')} → {os.get('year3', '')}", bold=True, color=MAROON)
    
    # Strategic Initiatives
    elif content_type == 'strategic_initiatives' and 'strategic_initiatives' in slide:
        for init in slide['strategic_initiatives']:
            add_paragraph_rtl(doc, f"• {init.get('initiative', '')}", bold=True, size=11)
            add_paragraph_rtl(doc, init.get('description', ''), size=10, color=GRAY)
            if 'components' in init:
                add_paragraph_rtl(doc, f"المكونات: {' | '.join(init['components'])}", size=10)
            if 'frequency' in init:
                add_paragraph_rtl(doc, f"التكرار: {init['frequency']}", size=10)
            if 'kpi' in init:
                add_paragraph_rtl(doc, f"المؤشر: {init['kpi']}", size=10, color=MAROON)
            doc.add_paragraph()
    
    # Channel Strategy
    elif content_type == 'channel_strategy' and 'channel_strategy' in slide:
        cs = slide['channel_strategy']
        if 'digital_channels' in cs:
            add_paragraph_rtl(doc, "القنوات الرقمية:", bold=True, color=MAROON)
            headers = ['القناة', 'المتابعون', 'المستهدف', 'الدور']
            rows = [[c.get('channel', ''), c.get('followers', c.get('subscribers', '')),
                    c.get('target', ''), c.get('role', '')] for c in cs['digital_channels']]
            create_table_rtl(doc, headers, rows)
        
        if 'traditional_channels' in cs:
            doc.add_paragraph()
            add_paragraph_rtl(doc, "القنوات التقليدية:", bold=True, color=MAROON)
            for tc in cs['traditional_channels']:
                add_bullet_rtl(doc, f"{tc.get('channel', '')}: {tc.get('type', '')} ({tc.get('frequency', '')})")
    
    # Campaign Calendar
    elif content_type == 'campaign_calendar' and 'campaign_calendar' in slide:
        cc = slide['campaign_calendar']
        if 'major_campaigns' in cc:
            add_paragraph_rtl(doc, "الحملات الرئيسية:", bold=True, color=MAROON)
            for camp in cc['major_campaigns']:
                add_paragraph_rtl(doc, f"• {camp.get('name', '')} - {camp.get('timing', '')}", bold=True, size=11)
                add_paragraph_rtl(doc, f"الميزانية: {camp.get('budget_allocation', '')}", size=10)
                if 'objectives' in camp:
                    add_paragraph_rtl(doc, f"الأهداف: {' | '.join(camp['objectives'])}", size=10)
                if 'kpis' in camp:
                    add_paragraph_rtl(doc, f"المؤشرات: {' | '.join(camp['kpis'])}", size=10, color=MAROON)
        
        if 'always_on' in cc:
            doc.add_paragraph()
            ao = cc['always_on']
            add_paragraph_rtl(doc, f"الحملات المستمرة ({ao.get('budget_allocation', '')})", bold=True, color=MAROON)
            add_paragraph_rtl(doc, ao.get('description', ''), size=11)
    
    # Monitoring System
    elif content_type == 'monitoring_system' and 'monitoring_system' in slide:
        ms = slide['monitoring_system']
        if 'components' in ms:
            for comp in ms['components']:
                add_paragraph_rtl(doc, f"• {comp.get('name', '')}", bold=True, size=11)
                if 'tools' in comp:
                    add_paragraph_rtl(doc, f"الأدوات: {' | '.join(comp['tools'])}", size=10)
                if 'outputs' in comp:
                    add_paragraph_rtl(doc, f"المخرجات: {' | '.join(comp['outputs'])}", size=10)
        
        if 'dashboards' in ms:
            doc.add_paragraph()
            add_paragraph_rtl(doc, "لوحات المتابعة:", bold=True, color=MAROON)
            headers = ['اللوحة', 'الجمهور', 'التكرار']
            rows = [[d.get('name', ''), d.get('audience', ''), d.get('frequency', '')] for d in ms['dashboards']]
            create_table_rtl(doc, headers, rows)
    
    # Crisis Framework
    elif content_type == 'crisis_framework' and 'crisis_framework' in slide:
        cf = slide['crisis_framework']
        if 'levels' in cf:
            add_paragraph_rtl(doc, "مستويات الأزمات:", bold=True, color=MAROON)
            for level in cf['levels']:
                add_paragraph_rtl(doc, f"• {level.get('level', '')}", bold=True, size=11)
                add_paragraph_rtl(doc, level.get('description', ''), size=10, color=GRAY)
                add_paragraph_rtl(doc, f"زمن الاستجابة: {level.get('response_time', '')} | صانع القرار: {level.get('decision_maker', '')}", size=10)
        
        if 'crisis_team' in cf:
            doc.add_paragraph()
            add_paragraph_rtl(doc, f"فريق الأزمات: {' | '.join(cf['crisis_team'])}", bold=True)
    
    # Internal Comms
    elif content_type == 'internal_comms' and 'internal_comms' in slide:
        ic = slide['internal_comms']
        if 'vision' in ic:
            add_paragraph_rtl(doc, f"الرؤية: {ic['vision']}", bold=True, color=MAROON)
        
        if 'target_audience' in ic:
            doc.add_paragraph()
            ta = ic['target_audience']
            add_paragraph_rtl(doc, f"إجمالي الموظفين: {ta.get('total_employees', '')}", size=11)
            if 'breakdown' in ta:
                headers = ['الشريحة', 'العدد', 'الاحتياجات']
                rows = [[b.get('segment', ''), str(b.get('count', '')), b.get('needs', '')] for b in ta['breakdown']]
                create_table_rtl(doc, headers, rows)
        
        if 'channels' in ic:
            doc.add_paragraph()
            add_paragraph_rtl(doc, "قنوات التواصل الداخلي:", bold=True, color=MAROON)
            for ch in ic['channels']:
                add_paragraph_rtl(doc, f"• {ch.get('channel', '')}", bold=True, size=11)
                add_paragraph_rtl(doc, ch.get('purpose', ''), size=10, color=GRAY)
        
        if 'kpis' in ic:
            doc.add_paragraph()
            add_paragraph_rtl(doc, "مؤشرات الأداء:", bold=True, color=MAROON)
            for kpi_name, kpi_val in ic['kpis'].items():
                if isinstance(kpi_val, dict):
                    add_bullet_rtl(doc, f"{kpi_name}: {kpi_val.get('current', '')} → {kpi_val.get('target', '')}")
    
    # Engagement Program
    elif content_type == 'engagement_program' and 'engagement_program' in slide:
        ep = slide['engagement_program']
        
        if 'current_state' in ep and 'target_state' in ep:
            add_paragraph_rtl(doc, "الوضع الحالي vs المستهدف:", bold=True, color=MAROON)
            headers = ['المؤشر', 'الحالي', 'المستهدف']
            cs = ep['current_state']
            ts = ep['target_state']
            rows = [
                ['مشاركة الموظفين', cs.get('engagement_score', ''), ts.get('engagement_score', '')],
                ['الرضا الوظيفي', cs.get('satisfaction_score', ''), ts.get('satisfaction_score', '')],
                ['معدل الدوران', cs.get('turnover_rate', ''), ts.get('turnover_rate', '')],
                ['eNPS', cs.get('enps', ''), ts.get('enps', '')]
            ]
            create_table_rtl(doc, headers, rows)
        
        if 'initiatives' in ep:
            doc.add_paragraph()
            add_paragraph_rtl(doc, "المبادرات:", bold=True, color=MAROON)
            for init in ep['initiatives']:
                add_paragraph_rtl(doc, f"• {init.get('name', '')}", bold=True, size=11)
                add_paragraph_rtl(doc, init.get('description', ''), size=10, color=GRAY)
                if 'budget' in init:
                    add_paragraph_rtl(doc, f"الميزانية: {init['budget']}", size=10, color=MAROON)
    
    # Ambassador Program
    elif content_type == 'ambassador_program' and 'ambassador_program' in slide:
        ap = slide['ambassador_program']
        if 'vision' in ap:
            add_paragraph_rtl(doc, f"الرؤية: {ap['vision']}", bold=True, color=MAROON)
        
        if 'program_structure' in ap:
            ps = ap['program_structure']
            doc.add_paragraph()
            add_paragraph_rtl(doc, f"اسم البرنامج: {ps.get('name', '')}", size=11)
            if 'target_ambassadors' in ps:
                ta = ps['target_ambassadors']
                add_paragraph_rtl(doc, f"المستهدف: السنة 1: {ta.get('year1', '')} | السنة 2: {ta.get('year2', '')} | السنة 3: {ta.get('year3', '')}", size=11)
            if 'selection_criteria' in ps:
                add_paragraph_rtl(doc, "معايير الاختيار:", bold=True)
                for crit in ps['selection_criteria']:
                    add_bullet_rtl(doc, crit)
        
        if 'training_program' in ap:
            doc.add_paragraph()
            add_paragraph_rtl(doc, "البرنامج التدريبي:", bold=True, color=MAROON)
            headers = ['الوحدة', 'المدة', 'المحتوى']
            rows = [[m.get('module', ''), m.get('duration', ''), ' | '.join(m.get('content', []))] 
                    for m in ap['training_program']]
            create_table_rtl(doc, headers, rows)
        
        if 'ambassador_benefits' in ap:
            doc.add_paragraph()
            add_paragraph_rtl(doc, "مزايا السفراء:", bold=True, color=MAROON)
            for benefit in ap['ambassador_benefits']:
                add_bullet_rtl(doc, benefit)
    
    # Leadership Comms
    elif content_type == 'leadership_comms' and 'leadership_comms' in slide:
        lc = slide['leadership_comms']
        if 'objectives' in lc:
            add_paragraph_rtl(doc, "الأهداف:", bold=True, color=MAROON)
            for obj in lc['objectives']:
                add_bullet_rtl(doc, obj)
        
        if 'leadership_visibility' in lc:
            lv = lc['leadership_visibility']
            if 'internal' in lv:
                doc.add_paragraph()
                add_paragraph_rtl(doc, "الحضور الداخلي:", bold=True, color=MAROON)
                for act in lv['internal']:
                    add_paragraph_rtl(doc, f"• {act.get('activity', '')} ({act.get('frequency', '')})", bold=True, size=11)
                    add_paragraph_rtl(doc, act.get('purpose', ''), size=10, color=GRAY)
            
            if 'external' in lv:
                doc.add_paragraph()
                add_paragraph_rtl(doc, "الحضور الخارجي:", bold=True, color=MAROON)
                for act in lv['external']:
                    add_paragraph_rtl(doc, f"• {act.get('activity', '')} ({act.get('frequency', '')})", bold=True, size=11)
                    add_paragraph_rtl(doc, act.get('purpose', ''), size=10, color=GRAY)
    
    # Knowledge Management
    elif content_type == 'knowledge_management' and 'knowledge_management' in slide:
        km = slide['knowledge_management']
        if 'objectives' in km:
            add_paragraph_rtl(doc, "الأهداف:", bold=True, color=MAROON)
            for obj in km['objectives']:
                add_bullet_rtl(doc, obj)
        
        if 'components' in km:
            doc.add_paragraph()
            add_paragraph_rtl(doc, "المكونات:", bold=True, color=MAROON)
            for comp in km['components']:
                add_paragraph_rtl(doc, f"• {comp.get('name', '')}", bold=True, size=11)
                add_paragraph_rtl(doc, comp.get('description', ''), size=10, color=GRAY)
        
        if 'knowledge_sharing_activities' in km:
            doc.add_paragraph()
            add_paragraph_rtl(doc, "أنشطة مشاركة المعرفة:", bold=True, color=MAROON)
            headers = ['النشاط', 'التكرار', 'الشكل']
            rows = [[a.get('activity', ''), a.get('frequency', ''), a.get('format', '')] 
                    for a in km['knowledge_sharing_activities']]
            create_table_rtl(doc, headers, rows)
    
    # Culture Building
    elif content_type == 'culture_building' and 'culture_building' in slide:
        cb = slide['culture_building']
        if 'core_values' in cb:
            add_paragraph_rtl(doc, "القيم الأساسية:", bold=True, color=MAROON)
            for val in cb['core_values']:
                add_paragraph_rtl(doc, f"• {val.get('value', '')}", bold=True, size=11)
                add_paragraph_rtl(doc, val.get('definition', ''), size=10, color=GRAY)
                if 'behaviors' in val:
                    add_paragraph_rtl(doc, f"السلوكيات: {' | '.join(val['behaviors'])}", size=10)
        
        if 'culture_initiatives' in cb:
            doc.add_paragraph()
            add_paragraph_rtl(doc, "مبادرات الثقافة:", bold=True, color=MAROON)
            for init in cb['culture_initiatives']:
                add_paragraph_rtl(doc, f"• {init.get('name', '')}", bold=True, size=11)
                add_paragraph_rtl(doc, init.get('description', ''), size=10, color=GRAY)
        
        if 'culture_metrics' in cb:
            doc.add_paragraph()
            add_paragraph_rtl(doc, "مؤشرات الثقافة:", bold=True, color=MAROON)
            for metric, values in cb['culture_metrics'].items():
                if isinstance(values, dict):
                    add_bullet_rtl(doc, f"{metric}: {values.get('current', '')} → {values.get('target', '')}")
    
    # Integration Matrix
    elif content_type == 'integration_matrix' and 'integration_matrix' in slide:
        im = slide['integration_matrix']
        if 'description' in im:
            add_paragraph_rtl(doc, im['description'], size=11, color=GRAY)
        
        if 'connections' in im:
            doc.add_paragraph()
            add_paragraph_rtl(doc, "روابط التكامل:", bold=True, color=MAROON)
            for conn in im['connections']:
                add_paragraph_rtl(doc, f"• من {conn.get('from', '')} إلى {conn.get('to', '')}", bold=True, size=11)
                add_paragraph_rtl(doc, conn.get('relationship', ''), size=10, color=GRAY)
        
        if 'sync_mechanisms' in im:
            doc.add_paragraph()
            add_paragraph_rtl(doc, "آليات التنسيق:", bold=True, color=MAROON)
            for mech in im['sync_mechanisms']:
                add_paragraph_rtl(doc, f"• {mech.get('mechanism', '')}", bold=True, size=11)
                add_paragraph_rtl(doc, mech.get('purpose', ''), size=10, color=GRAY)
    
    # KPI Dashboard
    elif content_type == 'kpi_dashboard' and 'kpi_dashboard' in slide:
        kd = slide['kpi_dashboard']
        
        if 'strategic_kpis' in kd:
            add_paragraph_rtl(doc, "المؤشرات الاستراتيجية:", bold=True, color=MAROON)
            headers = ['المؤشر', 'الحالي', 'السنة 1', 'السنة 3']
            rows = [[k.get('kpi', ''), str(k.get('current', '')), str(k.get('y1', '')), str(k.get('y3', ''))] 
                    for k in kd['strategic_kpis']]
            create_table_rtl(doc, headers, rows)
        
        if 'operational_kpis' in kd:
            doc.add_paragraph()
            add_paragraph_rtl(doc, "المؤشرات التشغيلية:", bold=True, color=MAROON)
            headers = ['المؤشر', 'الحالي', 'السنة 1', 'السنة 3']
            rows = [[k.get('kpi', ''), str(k.get('current', '')), str(k.get('y1', '')), str(k.get('y3', ''))] 
                    for k in kd['operational_kpis']]
            create_table_rtl(doc, headers, rows)
        
        if 'internal_kpis' in kd:
            doc.add_paragraph()
            add_paragraph_rtl(doc, "المؤشرات الداخلية:", bold=True, color=MAROON)
            headers = ['المؤشر', 'الحالي', 'السنة 1', 'السنة 3']
            rows = [[k.get('kpi', ''), str(k.get('current', '')), str(k.get('y1', '')), str(k.get('y3', ''))] 
                    for k in kd['internal_kpis']]
            create_table_rtl(doc, headers, rows)
    
    # Roadmap
    elif content_type == 'roadmap' and 'roadmap' in slide:
        rm = slide['roadmap']
        if 'phases' in rm:
            for phase in rm['phases']:
                add_paragraph_rtl(doc, f"• {phase.get('phase', '')} ({phase.get('duration', '')})", bold=True, size=12, color=MAROON)
                
                if 'track_activities' in phase:
                    ta = phase['track_activities']
                    if 'strategic' in ta:
                        add_paragraph_rtl(doc, "الاستراتيجي:", bold=True, size=10)
                        for act in ta['strategic']:
                            add_bullet_rtl(doc, act, level=1)
                    if 'operational' in ta:
                        add_paragraph_rtl(doc, "التشغيلي:", bold=True, size=10)
                        for act in ta['operational']:
                            add_bullet_rtl(doc, act, level=1)
                    if 'internal' in ta:
                        add_paragraph_rtl(doc, "الداخلي:", bold=True, size=10)
                        for act in ta['internal']:
                            add_bullet_rtl(doc, act, level=1)
                
                if 'key_deliverables' in phase:
                    add_paragraph_rtl(doc, f"المخرجات: {' | '.join(phase['key_deliverables'])}", size=10, color=GOLD)
                doc.add_paragraph()
    
    # Team Structure
    elif content_type == 'team_structure' and 'team_structure' in slide:
        ts = slide['team_structure']
        if 'proposed_structure' in ts:
            ps = ts['proposed_structure']
            if 'head' in ps:
                head = ps['head']
                add_paragraph_rtl(doc, f"رئيس الفريق: {head.get('title', '')}", bold=True, color=MAROON)
                add_paragraph_rtl(doc, f"يرفع تقاريره إلى: {head.get('reports_to', '')}", size=11)
            
            if 'units' in ps:
                doc.add_paragraph()
                add_paragraph_rtl(doc, "الوحدات:", bold=True, color=MAROON)
                headers = ['الوحدة', 'الرئيس', 'الحجم', 'المسؤوليات']
                rows = [[u.get('name', ''), u.get('head', ''), str(u.get('team_size', '')), 
                        ' | '.join(u.get('responsibilities', []))] for u in ps['units']]
                create_table_rtl(doc, headers, rows)
        
        if 'total_team' in ts:
            doc.add_paragraph()
            add_paragraph_rtl(doc, f"إجمالي الفريق: {ts['total_team']} شخص", bold=True, color=MAROON)
        
        if 'external_support' in ts:
            doc.add_paragraph()
            add_paragraph_rtl(doc, "الدعم الخارجي:", bold=True, color=MAROON)
            for ext in ts['external_support']:
                add_bullet_rtl(doc, f"{ext.get('role', '')}: {ext.get('scope', '')}")
    
    # Budget Overview
    elif content_type == 'budget_overview' and 'budget_overview' in slide:
        bo = slide['budget_overview']
        if 'note' in bo:
            add_paragraph_rtl(doc, f"ملاحظة: {bo['note']}", size=10, color=GRAY)
        
        if 'annual_distribution' in bo:
            doc.add_paragraph()
            add_paragraph_rtl(doc, "التوزيع السنوي:", bold=True, color=MAROON)
            ad = bo['annual_distribution']
            for track_name, track_data in ad.items():
                if isinstance(track_data, dict):
                    add_paragraph_rtl(doc, f"• {track_name}: {track_data.get('percentage', '')}", bold=True, size=11)
                    if 'main_items' in track_data:
                        add_paragraph_rtl(doc, f"البنود: {' | '.join(track_data['main_items'])}", size=10)
        
        if 'roi_projections' in bo:
            doc.add_paragraph()
            add_paragraph_rtl(doc, "توقعات العائد:", bold=True, color=MAROON)
            rp = bo['roi_projections']
            add_paragraph_rtl(doc, f"السنة 1: {rp.get('year1', '')}", size=11)
            add_paragraph_rtl(doc, f"السنة 2: {rp.get('year2', '')}", size=11)
            add_paragraph_rtl(doc, f"السنة 3: {rp.get('year3', '')}", size=11)
    
    # Differentiators
    elif content_type == 'differentiators' and 'differentiators' in slide:
        diff = slide['differentiators']
        if 'approach_differentiators' in diff:
            add_paragraph_rtl(doc, "مميزات المنهج:", bold=True, color=MAROON)
            for d in diff['approach_differentiators']:
                add_paragraph_rtl(doc, f"• {d.get('title', '')}", bold=True, size=11)
                add_paragraph_rtl(doc, d.get('description', ''), size=10, color=GRAY)
                if 'value' in d:
                    add_paragraph_rtl(doc, f"القيمة: {d['value']}", size=10, color=GOLD)
        
        if 'team_differentiators' in diff:
            doc.add_paragraph()
            add_paragraph_rtl(doc, "مميزات الفريق:", bold=True, color=MAROON)
            for d in diff['team_differentiators']:
                add_paragraph_rtl(doc, f"• {d.get('title', '')}", bold=True, size=11)
                add_paragraph_rtl(doc, d.get('description', ''), size=10, color=GRAY)
    
    # Success Stories
    elif content_type == 'success_stories' and 'success_stories' in slide:
        for story in slide['success_stories']:
            add_paragraph_rtl(doc, f"• {story.get('client', '')} - {story.get('sector', '')}", bold=True, size=11, color=MAROON)
            add_paragraph_rtl(doc, f"التحدي: {story.get('challenge', '')}", size=10)
            add_paragraph_rtl(doc, f"الحل: {story.get('solution', '')}", size=10)
            if 'results' in story:
                add_paragraph_rtl(doc, "النتائج:", bold=True, size=10)
                for result in story['results']:
                    add_bullet_rtl(doc, result, level=1)
            doc.add_paragraph()
    
    add_page_break(doc)


def render_closing(doc, slide):
    """Render closing slide"""
    doc.add_paragraph()
    doc.add_paragraph()
    
    t = add_heading_rtl(doc, slide.get('title', ''), level=1, color=MAROON)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if 'message' in slide:
        m = add_paragraph_rtl(doc, slide['message'], size=14, color=GOLD)
        m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if 'next_steps' in slide:
        doc.add_paragraph()
        add_paragraph_rtl(doc, "الخطوات القادمة:", bold=True, color=MAROON)
        for step in slide['next_steps']:
            if isinstance(step, dict):
                add_paragraph_rtl(doc, f"• {step.get('step', '')}", bold=True, size=11)
                add_paragraph_rtl(doc, step.get('description', ''), size=10, color=GRAY)
            else:
                add_bullet_rtl(doc, step)
    
    if 'closing_message' in slide:
        doc.add_paragraph()
        doc.add_paragraph()
        cm = add_paragraph_rtl(doc, slide['closing_message'], size=12)
        cm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if 'contact' in slide:
        doc.add_paragraph()
        doc.add_paragraph()
        contact = slide['contact']
        if 'message' in contact:
            c = add_paragraph_rtl(doc, contact['message'], size=12, color=GRAY)
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if 'cta' in contact:
            cta = add_paragraph_rtl(doc, contact['cta'], bold=True, size=16, color=MAROON)
            cta.alignment = WD_ALIGN_PARAGRAPH.CENTER


def render_introduction(doc, slide):
    """Render introduction slide"""
    add_heading_rtl(doc, slide.get('title', ''), level=1)
    if 'content' in slide:
        # Split by paragraphs
        paragraphs = slide['content'].split('\n\n')
        for para in paragraphs:
            if para.strip():
                add_paragraph_rtl(doc, para.strip(), size=12)
    add_page_break(doc)


def render_context_analysis(doc, slide):
    """Render context analysis slide"""
    add_heading_rtl(doc, slide.get('title', ''), level=1)
    
    if 'introduction' in slide:
        add_paragraph_rtl(doc, slide['introduction'], size=11, color=GRAY)
    
    if 'challenges' in slide:
        doc.add_paragraph()
        add_paragraph_rtl(doc, "التحديات:", bold=True, color=MAROON, size=14)
        for ch in slide['challenges']:
            add_paragraph_rtl(doc, f"• {ch.get('challenge', '')}", bold=True, size=11)
            add_paragraph_rtl(doc, ch.get('description', ''), size=10, color=GRAY)
            doc.add_paragraph()
    
    if 'opportunities' in slide:
        doc.add_paragraph()
        add_paragraph_rtl(doc, "الفرص:", bold=True, color=GOLD, size=14)
        for op in slide['opportunities']:
            add_paragraph_rtl(doc, f"• {op.get('opportunity', '')}", bold=True, size=11)
            add_paragraph_rtl(doc, op.get('description', ''), size=10, color=GRAY)
            doc.add_paragraph()
    
    add_page_break(doc)


def render_methodology_overview(doc, slide):
    """Render methodology overview slide"""
    add_heading_rtl(doc, slide.get('title', ''), level=1)
    
    if 'introduction' in slide:
        add_paragraph_rtl(doc, slide['introduction'], size=11)
    
    if 'tracks' in slide:
        doc.add_paragraph()
        for track in slide['tracks']:
            icon = track.get('icon', '')
            name = track.get('name', '')
            add_paragraph_rtl(doc, f"{icon} {name}", bold=True, color=MAROON, size=14)
            
            if 'philosophy' in track:
                add_paragraph_rtl(doc, track['philosophy'], size=11)
            
            if 'components' in track:
                doc.add_paragraph()
                add_paragraph_rtl(doc, "المكونات:", bold=True, size=11)
                for comp in track['components']:
                    add_bullet_rtl(doc, comp)
            
            if 'outcomes' in track:
                doc.add_paragraph()
                add_paragraph_rtl(doc, "المخرجات:", bold=True, size=11, color=GOLD)
                for out in track['outcomes']:
                    add_bullet_rtl(doc, out)
            
            doc.add_paragraph()
    
    if 'integration_note' in slide:
        doc.add_paragraph()
        add_paragraph_rtl(doc, slide['integration_note'], size=11, color=GRAY)
    
    add_page_break(doc)


def render_detailed_component(doc, slide):
    """Render detailed component slide"""
    title = slide.get('title', '')
    number = slide.get('component_number', '')
    add_heading_rtl(doc, f"{number} - {title}" if number else title, level=1)
    
    if 'introduction' in slide:
        add_paragraph_rtl(doc, slide['introduction'], size=11)
    
    # Deliverables section
    if 'deliverables' in slide:
        doc.add_paragraph()
        add_paragraph_rtl(doc, "المخرجات:", bold=True, color=MAROON, size=14)
        for key, val in slide['deliverables'].items():
            if isinstance(val, dict):
                add_paragraph_rtl(doc, f"• {val.get('title', key)}", bold=True, size=12)
                if 'description' in val:
                    add_paragraph_rtl(doc, val['description'], size=11)
                if 'approach' in val:
                    add_paragraph_rtl(doc, f"المنهج: {val['approach']}", size=10, color=GRAY)
                if 'dimensions' in val:
                    for dim_key, dim_val in val['dimensions'].items():
                        add_paragraph_rtl(doc, f"- {dim_key}: {dim_val}", size=10)
                if 'example_pillars' in val:
                    doc.add_paragraph()
                    for pillar in val['example_pillars']:
                        add_paragraph_rtl(doc, f"  • {pillar.get('pillar', '')}", bold=True, size=11)
                        add_paragraph_rtl(doc, f"    {pillar.get('core_message', '')}", size=10)
                doc.add_paragraph()
    
    # Stakeholder circles
    if 'stakeholder_circles' in slide:
        doc.add_paragraph()
        for circle in slide['stakeholder_circles']:
            add_paragraph_rtl(doc, f"• {circle.get('circle', '')}", bold=True, size=12, color=MAROON)
            add_paragraph_rtl(doc, circle.get('description', ''), size=11)
            
            if 'stakeholders' in circle:
                doc.add_paragraph()
                for sh in circle['stakeholders']:
                    if isinstance(sh, dict):
                        add_paragraph_rtl(doc, f"  - {sh.get('segment', '')}", bold=True, size=10)
                        if 'description' in sh:
                            add_paragraph_rtl(doc, f"    {sh.get('description', '')}", size=10, color=GRAY)
                        if 'estimated_count' in sh:
                            add_paragraph_rtl(doc, f"    العدد: {sh.get('estimated_count', '')}", size=10)
                        if 'primary_needs' in sh:
                            add_paragraph_rtl(doc, f"    الاحتياجات: {sh.get('primary_needs', '')}", size=10)
                        if 'communication_preferences' in sh:
                            add_paragraph_rtl(doc, f"    تفضيلات التواصل: {sh.get('communication_preferences', '')}", size=10)
                        if 'relationship_nature' in sh:
                            add_paragraph_rtl(doc, f"    طبيعة العلاقة: {sh.get('relationship_nature', '')}", size=10)
                        if 'examples' in sh:
                            add_paragraph_rtl(doc, f"    أمثلة: {sh.get('examples', '')}", size=10)
            doc.add_paragraph()
    
    # Measurement model
    if 'measurement_model' in slide:
        mm = slide['measurement_model']
        add_paragraph_rtl(doc, f"النموذج: {mm.get('name', '')}", bold=True, color=MAROON)
        if 'description' in mm:
            add_paragraph_rtl(doc, mm['description'], size=11)
        
        if 'dimensions' in mm:
            doc.add_paragraph()
            for dim in mm['dimensions']:
                add_paragraph_rtl(doc, f"• {dim.get('dimension', '')} ({dim.get('weight', '')})", bold=True, size=11)
                add_paragraph_rtl(doc, dim.get('description', ''), size=10, color=GRAY)
                if 'indicators' in dim:
                    for ind in dim['indicators']:
                        add_bullet_rtl(doc, ind, level=1)
    
    # Measurement methodology
    if 'measurement_methodology' in slide:
        mm = slide['measurement_methodology']
        doc.add_paragraph()
        add_paragraph_rtl(doc, mm.get('title', 'منهجية القياس'), bold=True, color=MAROON)
        if 'components' in mm:
            for comp in mm['components']:
                add_paragraph_rtl(doc, f"• {comp.get('component', '')}", bold=True, size=11)
                add_paragraph_rtl(doc, f"  التكرار: {comp.get('frequency', '')}", size=10)
                add_paragraph_rtl(doc, f"  {comp.get('description', '')}", size=10, color=GRAY)
    
    # Reporting
    if 'reporting' in slide:
        rp = slide['reporting']
        doc.add_paragraph()
        add_paragraph_rtl(doc, rp.get('title', 'التقارير'), bold=True, color=MAROON)
        if 'outputs' in rp:
            for out in rp['outputs']:
                add_paragraph_rtl(doc, f"• {out.get('name', '')}", bold=True, size=11)
                add_paragraph_rtl(doc, f"  الجمهور: {out.get('audience', '')} | التكرار: {out.get('frequency', '')}", size=10)
                add_paragraph_rtl(doc, f"  {out.get('content', '')}", size=10, color=GRAY)
    
    # Content pillars
    if 'content_pillars' in slide:
        doc.add_paragraph()
        add_paragraph_rtl(doc, "ركائز المحتوى:", bold=True, color=MAROON, size=14)
        for pillar in slide['content_pillars']:
            add_paragraph_rtl(doc, f"• {pillar.get('pillar', '')}", bold=True, size=12)
            add_paragraph_rtl(doc, pillar.get('description', ''), size=11)
            if 'content_types' in pillar:
                for ct in pillar['content_types']:
                    add_bullet_rtl(doc, ct, level=1)
            if 'success_metrics' in pillar:
                add_paragraph_rtl(doc, f"  مؤشرات النجاح: {pillar['success_metrics']}", size=10, color=GOLD)
            doc.add_paragraph()
    
    # Content calendar
    if 'content_calendar' in slide:
        cc = slide['content_calendar']
        doc.add_paragraph()
        add_paragraph_rtl(doc, cc.get('title', 'تقويم المحتوى'), bold=True, color=MAROON)
        if 'description' in cc:
            add_paragraph_rtl(doc, cc['description'], size=11)
        if 'key_moments' in cc:
            add_paragraph_rtl(doc, "المناسبات الرئيسية:", bold=True, size=11)
            for moment in cc['key_moments']:
                add_bullet_rtl(doc, moment)
    
    # Quality standards
    if 'quality_standards' in slide:
        qs = slide['quality_standards']
        doc.add_paragraph()
        add_paragraph_rtl(doc, qs.get('title', 'معايير الجودة'), bold=True, color=MAROON)
        if 'standards' in qs:
            for std in qs['standards']:
                add_bullet_rtl(doc, std)
    
    # Digital channels
    if 'digital_channels' in slide:
        dc = slide['digital_channels']
        doc.add_paragraph()
        add_paragraph_rtl(doc, dc.get('title', 'القنوات الرقمية'), bold=True, color=MAROON, size=14)
        if 'philosophy' in dc:
            add_paragraph_rtl(doc, dc['philosophy'], size=11)
        if 'channels' in dc:
            for ch in dc['channels']:
                add_paragraph_rtl(doc, f"• {ch.get('platform', '')}", bold=True, size=12)
                add_paragraph_rtl(doc, f"  الدور: {ch.get('role', '')}", size=10)
                add_paragraph_rtl(doc, f"  استراتيجية المحتوى: {ch.get('content_strategy', '')}", size=10)
                add_paragraph_rtl(doc, f"  الإيقاع: {ch.get('posting_rhythm', '')}", size=10, color=GRAY)
                doc.add_paragraph()
    
    # Traditional channels
    if 'traditional_channels' in slide:
        tc = slide['traditional_channels']
        doc.add_paragraph()
        add_paragraph_rtl(doc, tc.get('title', 'القنوات التقليدية'), bold=True, color=MAROON, size=14)
        if 'philosophy' in tc:
            add_paragraph_rtl(doc, tc['philosophy'], size=11)
        if 'channels' in tc:
            for ch in tc['channels']:
                add_paragraph_rtl(doc, f"• {ch.get('channel', '')}", bold=True, size=11)
                add_paragraph_rtl(doc, f"  المنهج: {ch.get('approach', '')}", size=10)
                if 'activities' in ch:
                    add_paragraph_rtl(doc, f"  الأنشطة: {ch['activities']}", size=10, color=GRAY)
    
    # Channel integration
    if 'channel_integration' in slide:
        ci = slide['channel_integration']
        doc.add_paragraph()
        add_paragraph_rtl(doc, ci.get('title', 'تكامل القنوات'), bold=True, color=MAROON)
        add_paragraph_rtl(doc, ci.get('description', ''), size=11)
    
    # Campaign types
    if 'campaign_types' in slide:
        doc.add_paragraph()
        for ct in slide['campaign_types']:
            add_paragraph_rtl(doc, f"• {ct.get('type', '')}", bold=True, size=12, color=MAROON)
            add_paragraph_rtl(doc, ct.get('description', ''), size=11)
            
            if 'examples' in ct:
                for ex in ct['examples']:
                    add_paragraph_rtl(doc, f"  - {ex.get('campaign', '')} ({ex.get('timing', '')})", bold=True, size=11)
                    if 'objectives' in ex:
                        add_paragraph_rtl(doc, "    الأهداف:", bold=True, size=10)
                        for obj in ex['objectives']:
                            add_bullet_rtl(doc, obj, level=2)
                    if 'key_messages' in ex:
                        add_paragraph_rtl(doc, "    الرسائل الرئيسية:", bold=True, size=10)
                        for msg in ex['key_messages']:
                            add_bullet_rtl(doc, msg, level=2)
                    if 'activities' in ex:
                        add_paragraph_rtl(doc, "    الأنشطة:", bold=True, size=10)
                        for act in ex['activities']:
                            add_bullet_rtl(doc, act, level=2)
            
            if 'campaigns' in ct:
                for camp in ct['campaigns']:
                    add_paragraph_rtl(doc, f"  - {camp.get('name', '')}", bold=True, size=11)
                    add_paragraph_rtl(doc, f"    الهدف: {camp.get('objective', '')}", size=10)
                    add_paragraph_rtl(doc, f"    المنهج: {camp.get('approach', '')}", size=10, color=GRAY)
            doc.add_paragraph()
    
    # Campaign methodology
    if 'campaign_methodology' in slide:
        cm = slide['campaign_methodology']
        doc.add_paragraph()
        add_paragraph_rtl(doc, cm.get('title', ''), bold=True, color=MAROON)
        if 'phases' in cm:
            for phase in cm['phases']:
                add_paragraph_rtl(doc, f"• {phase.get('phase', '')}", bold=True, size=11)
                if 'activities' in phase:
                    for act in phase['activities']:
                        add_bullet_rtl(doc, act, level=1)
    
    # Monitoring components
    if 'monitoring_components' in slide:
        doc.add_paragraph()
        add_paragraph_rtl(doc, "مكونات الرصد:", bold=True, color=MAROON, size=14)
        for mc in slide['monitoring_components']:
            add_paragraph_rtl(doc, f"• {mc.get('component', '')}", bold=True, size=12)
            add_paragraph_rtl(doc, f"  النطاق: {mc.get('scope', '')}", size=10)
            if 'capabilities' in mc:
                add_paragraph_rtl(doc, "  القدرات:", bold=True, size=10)
                for cap in mc['capabilities']:
                    add_bullet_rtl(doc, cap, level=1)
            if 'outputs' in mc:
                add_paragraph_rtl(doc, f"  المخرجات: {' | '.join(mc['outputs'])}", size=10, color=GOLD)
            doc.add_paragraph()
    
    # Analysis framework
    if 'analysis_framework' in slide:
        af = slide['analysis_framework']
        doc.add_paragraph()
        add_paragraph_rtl(doc, af.get('title', 'إطار التحليل'), bold=True, color=MAROON)
        add_paragraph_rtl(doc, af.get('description', ''), size=11)
        if 'analysis_types' in af:
            for at in af['analysis_types']:
                add_paragraph_rtl(doc, f"• {at.get('type', '')}", bold=True, size=11)
                add_paragraph_rtl(doc, f"  السؤال: {at.get('question', '')}", size=10)
                add_paragraph_rtl(doc, f"  المخرجات: {at.get('outputs', '')}", size=10, color=GRAY)
    
    # Dashboards
    if 'dashboards' in slide:
        db = slide['dashboards']
        doc.add_paragraph()
        add_paragraph_rtl(doc, db.get('title', 'لوحات المتابعة'), bold=True, color=MAROON)
        add_paragraph_rtl(doc, db.get('description', ''), size=11)
        if 'dashboard_types' in db:
            headers = ['اللوحة', 'الجمهور', 'المحتوى', 'التحديث']
            rows = [[d.get('name', ''), d.get('audience', ''), d.get('content', ''), d.get('update_frequency', '')] 
                    for d in db['dashboard_types']]
            create_table_rtl(doc, headers, rows)
    
    # Crisis levels
    if 'crisis_levels' in slide:
        cl = slide['crisis_levels']
        doc.add_paragraph()
        add_paragraph_rtl(doc, cl.get('title', 'مستويات الأزمات'), bold=True, color=MAROON, size=14)
        add_paragraph_rtl(doc, cl.get('description', ''), size=11)
        if 'levels' in cl:
            for level in cl['levels']:
                add_paragraph_rtl(doc, f"• {level.get('level', '')}", bold=True, size=12)
                add_paragraph_rtl(doc, f"  التعريف: {level.get('definition', '')}", size=10)
                add_paragraph_rtl(doc, f"  أمثلة: {level.get('examples', '')}", size=10, color=GRAY)
                add_paragraph_rtl(doc, f"  زمن الاستجابة: {level.get('response_time', '')} | صانع القرار: {level.get('decision_authority', '')}", size=10)
                if 'actions' in level:
                    add_paragraph_rtl(doc, "  الإجراءات:", bold=True, size=10)
                    for act in level['actions']:
                        add_bullet_rtl(doc, act, level=1)
                doc.add_paragraph()
    
    # Crisis team
    if 'crisis_team' in slide:
        ct = slide['crisis_team']
        doc.add_paragraph()
        add_paragraph_rtl(doc, ct.get('title', 'فريق الأزمات'), bold=True, color=MAROON)
        if 'composition' in ct:
            for member in ct['composition']:
                add_paragraph_rtl(doc, f"• {member.get('role', '')}", bold=True, size=11)
                add_paragraph_rtl(doc, f"  المنصب: {member.get('typical_position', '')}", size=10)
                add_paragraph_rtl(doc, f"  المسؤوليات: {member.get('responsibilities', '')}", size=10, color=GRAY)
    
    # Preparedness
    if 'preparedness' in slide:
        prep = slide['preparedness']
        doc.add_paragraph()
        add_paragraph_rtl(doc, prep.get('title', 'الاستعداد'), bold=True, color=MAROON)
        if 'components' in prep:
            for comp in prep['components']:
                add_paragraph_rtl(doc, f"• {comp.get('component', '')}", bold=True, size=11)
                add_paragraph_rtl(doc, comp.get('description', ''), size=10, color=GRAY)
    
    # Methodology
    if 'methodology' in slide:
        meth = slide['methodology']
        doc.add_paragraph()
        add_paragraph_rtl(doc, meth.get('title', 'المنهجية'), bold=True, color=MAROON)
        if 'phases' in meth:
            for phase in meth['phases']:
                add_paragraph_rtl(doc, f"• {phase.get('phase', '')}", bold=True, size=11)
                if 'activities' in phase:
                    for act in phase['activities']:
                        add_bullet_rtl(doc, act, level=1)
        if 'steps' in meth:
            for step in meth['steps']:
                add_bullet_rtl(doc, step)
    
    add_page_break(doc)


def render_implementation_plan(doc, slide):
    """Render implementation plan slide"""
    add_heading_rtl(doc, slide.get('title', ''), level=1)
    
    if 'introduction' in slide:
        add_paragraph_rtl(doc, slide['introduction'], size=11)
    
    if 'phases' in slide:
        doc.add_paragraph()
        for phase in slide['phases']:
            add_paragraph_rtl(doc, f"• {phase.get('phase', '')} ({phase.get('duration', '')})", bold=True, size=12, color=MAROON)
            
            if 'objectives' in phase:
                add_paragraph_rtl(doc, "الأهداف:", bold=True, size=11)
                for obj in phase['objectives']:
                    add_bullet_rtl(doc, obj, level=1)
            
            if 'key_activities' in phase:
                ka = phase['key_activities']
                if 'strategic_track' in ka:
                    add_paragraph_rtl(doc, "المسار الاستراتيجي:", bold=True, size=10)
                    for act in ka['strategic_track']:
                        add_bullet_rtl(doc, act, level=1)
                if 'operational_track' in ka:
                    add_paragraph_rtl(doc, "المسار التشغيلي:", bold=True, size=10)
                    for act in ka['operational_track']:
                        add_bullet_rtl(doc, act, level=1)
            
            if 'deliverables' in phase:
                add_paragraph_rtl(doc, f"المخرجات: {' | '.join(phase['deliverables'])}", size=10, color=GOLD)
            
            doc.add_paragraph()
    
    add_page_break(doc)


def render_team_approach(doc, slide):
    """Render team approach slide"""
    add_heading_rtl(doc, slide.get('title', ''), level=1)
    
    if 'introduction' in slide:
        add_paragraph_rtl(doc, slide['introduction'], size=11)
    
    if 'our_team' in slide:
        ot = slide['our_team']
        doc.add_paragraph()
        add_paragraph_rtl(doc, ot.get('title', 'فريقنا'), bold=True, color=MAROON, size=14)
        if 'composition' in ot:
            for member in ot['composition']:
                add_paragraph_rtl(doc, f"• {member.get('role', '')}", bold=True, size=11)
                add_paragraph_rtl(doc, member.get('responsibilities', ''), size=10, color=GRAY)
    
    if 'collaboration_model' in slide:
        cm = slide['collaboration_model']
        doc.add_paragraph()
        add_paragraph_rtl(doc, cm.get('title', 'نموذج التعاون'), bold=True, color=MAROON, size=14)
        if 'principles' in cm:
            for prin in cm['principles']:
                add_paragraph_rtl(doc, f"• {prin.get('principle', '')}", bold=True, size=11)
                add_paragraph_rtl(doc, prin.get('description', ''), size=10, color=GRAY)
        
        if 'communication_rhythm' in cm:
            doc.add_paragraph()
            add_paragraph_rtl(doc, "إيقاع التواصل:", bold=True, size=11)
            headers = ['الاجتماع', 'المشاركون', 'الغرض']
            rows = [[m.get('meeting', ''), m.get('participants', ''), m.get('purpose', '')] 
                    for m in cm['communication_rhythm']]
            create_table_rtl(doc, headers, rows)
    
    add_page_break(doc)


def render_why_us(doc, slide):
    """Render why us slide"""
    add_heading_rtl(doc, slide.get('title', ''), level=1)
    
    if 'introduction' in slide:
        add_paragraph_rtl(doc, slide['introduction'], size=11)
    
    if 'differentiators' in slide:
        doc.add_paragraph()
        for diff in slide['differentiators']:
            add_paragraph_rtl(doc, f"• {diff.get('title', '')}", bold=True, size=12, color=MAROON)
            add_paragraph_rtl(doc, diff.get('description', ''), size=11)
            doc.add_paragraph()
    
    add_page_break(doc)


def export_to_word(json_path, output_path):
    """Main export function"""
    # Load JSON
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Create document
    doc = Document()
    
    # Set document to RTL
    for section in doc.sections:
        sectPr = section._sectPr
        bidi = OxmlElement('w:bidi')
        sectPr.append(bidi)
    
    # Process slides
    slides = data.get('slides', [])
    
    for slide in slides:
        slide_type = slide.get('type', '')
        
        if slide_type == 'cover':
            render_cover(doc, slide)
        elif slide_type == 'introduction':
            render_introduction(doc, slide)
        elif slide_type == 'context_analysis':
            render_context_analysis(doc, slide)
        elif slide_type == 'methodology_overview':
            render_methodology_overview(doc, slide)
        elif slide_type == 'executive_summary':
            render_executive_summary(doc, slide)
        elif slide_type == 'framework_overview':
            render_framework_overview(doc, slide)
        elif slide_type == 'section_divider':
            render_section_divider(doc, slide)
        elif slide_type == 'detailed_component':
            render_detailed_component(doc, slide)
        elif slide_type == 'pillar_detail':
            render_pillar_detail(doc, slide)
        elif slide_type == 'implementation_plan':
            render_implementation_plan(doc, slide)
        elif slide_type == 'team_approach':
            render_team_approach(doc, slide)
        elif slide_type == 'why_us':
            render_why_us(doc, slide)
        elif slide_type == 'closing':
            render_closing(doc, slide)
    
    # Save document
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"✅ تم تصدير الملف بنجاح: {output_path}")
    return output_path


if __name__ == "__main__":
    json_path = "data/presentations/cda_dubai_pitch_v4.json"
    output_path = "exports/CDA_Dubai_Strategic_2026.docx"
    export_to_word(json_path, output_path)
